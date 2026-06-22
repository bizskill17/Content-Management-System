from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


DEFAULT_COURSE_ROOT = Path(r"D:\GST for the Layman")
MAX_IMAGE_WIDTH = Inches(6.5)


def clean_name(text: str) -> str:
    text = re.sub(r"^\d+\.\s*", "", text.strip())
    text = re.sub(r"^[A-Z]\.\s*", "", text.strip())
    return text.strip().lower()


def iter_course_docs(course_root: Path) -> list[Path]:
    docs: list[Path] = []

    intro = course_root / "Introduction.docx"
    if intro.exists():
        docs.append(intro)

    for section_dir in sorted(course_root.iterdir()):
        if not section_dir.is_dir() or section_dir.name == "Ignore":
            continue
        docs.extend(sorted(section_dir.glob("*.docx")))

    return docs


def build_image_map(featured_dir: Path) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for image_path in sorted(featured_dir.glob("*.png")):
        if image_path.name.startswith("test-"):
            continue
        mapping[clean_name(image_path.stem)] = image_path
    return mapping


def insert_paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def has_nearby_drawing(paragraphs, start_index: int, lookahead: int = 3) -> bool:
    end_index = min(len(paragraphs), start_index + lookahead + 1)
    for idx in range(start_index + 1, end_index):
        xml = paragraphs[idx]._element.xml
        if "pic:pic" in xml or "w:drawing" in xml:
            return True
    return False


def add_image_below_title(doc_path: Path, image_path: Path, dry_run: bool = False) -> bool:
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    title_index = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip():
            title_index = idx
            break

    if title_index is None:
        return False

    if has_nearby_drawing(paragraphs, title_index):
        return False

    title_paragraph = paragraphs[title_index]
    image_paragraph = insert_paragraph_after(title_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=MAX_IMAGE_WIDTH)

    if dry_run:
        return True

    doc.save(doc_path)
    return True


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Insert matching featured images below chapter titles in course Word files."
    )
    parser.add_argument("--course-root", type=Path, default=DEFAULT_COURSE_ROOT)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    course_root = args.course_root
    featured_dir = course_root / "Ignore" / "featured_images"

    if not course_root.exists():
        raise FileNotFoundError(f"Course folder not found: {course_root}")
    if not featured_dir.exists():
        raise FileNotFoundError(f"Featured images folder not found: {featured_dir}")

    image_map = build_image_map(featured_dir)
    docs = iter_course_docs(course_root)

    updated = 0
    skipped = 0

    for doc_path in docs:
        key = clean_name(doc_path.stem)
        image_path = image_map.get(key)
        if image_path is None:
            print(f"SKIP  no matching image: {doc_path.name}")
            skipped += 1
            continue

        changed = add_image_below_title(doc_path, image_path, dry_run=args.dry_run)
        if changed:
            action = "WOULD UPDATE" if args.dry_run else "UPDATED"
            print(f"{action}  {doc_path.name} <- {image_path.name}")
            updated += 1
        else:
            print(f"SKIP  already has nearby image or no title: {doc_path.name}")
            skipped += 1

    print()
    print(f"Updated: {updated}")
    print(f"Skipped: {skipped}")


if __name__ == "__main__":
    main()
